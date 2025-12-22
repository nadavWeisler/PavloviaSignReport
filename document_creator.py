"""Create Word documents from tabular data with text and image columns.

This module provides functions to generate individual documents per row
and a single summary document with a table. Supports STRING and IMAGE
column types defined in settings.ColumnType.
"""

import base64
import logging
import os
from io import BytesIO

from PIL import Image
from docx import Document
from docx.shared import Inches

from settings import ColumnType

logger = logging.getLogger(__name__)


def create_single_doc(row, folder_path, columns, document_settings):
    """
    Create a Word document for a row in the data frame and save it to the specified folder.
    Args:
        row (pd.Series): The row in the data frame.
        folder_path (str): The path to the folder where the document will be saved.
    """
    doc = Document()
    doc.add_heading(document_settings.title, 0)

    for column in columns:
        if column.column_type == ColumnType.STRING:
            doc.add_paragraph(f"{column.display_name}: {row[column.name]}")
        elif column.column_type == ColumnType.IMAGE:
            try:
                base64_signature = row[column.name]
                base64_signature = base64_signature.replace(
                    "data:image/png;base64,", ""
                )
                signature_bytes = base64.b64decode(base64_signature)
                img = Image.open(BytesIO(signature_bytes))

                temp_signature_path = "temp_signature.png"
                img.save(temp_signature_path)

                doc.add_picture(
                    temp_signature_path,
                    width=Inches(document_settings.image_width),
                    height=Inches(document_settings.image_height),
                )

                os.remove(temp_signature_path)
            except Exception as e:
                print(e)
                doc.add_paragraph(f"{column.display_name}: {row[column.name]}")

    doc.save(f"{folder_path}/{row[document_settings.id_col]}.docx")


def create_doc_with_table(df, folder_path, columns, document_settings):
    """
    Create a single Word document with a table containing all rows from the dataframe.
    Args:
        df (pd.DataFrame): The dataframe containing all rows.
        folder_path (str): The path to the folder where the document will be saved.
    """
    doc = Document()
    doc.add_heading(document_settings.title, 0)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, column in enumerate(columns):
        header_cells[idx].text = column.display_name

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for idx, column in enumerate(columns):
            if column.column_type == ColumnType.STRING:
                row_cells[idx].text = str(row[column.name])
            elif column.column_type == ColumnType.IMAGE:
                try:
                    base64_signature = row[column.name]
                    base64_signature = base64_signature.replace(
                        "data:image/png;base64,", ""
                    )
                    signature_bytes = base64.b64decode(base64_signature)
                    img = Image.open(BytesIO(signature_bytes))

                    temp_signature_path = "temp_signature.png"
                    img.save(temp_signature_path)

                    paragraph = row_cells[idx].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        temp_signature_path, width=Inches(2.0), height=Inches(2.0)
                    )

                    os.remove(temp_signature_path)
                except Exception as e:
                    print(e)
                    row_cells[idx].text = str(row[column.name])

    doc.save(f"{folder_path}/summary.docx")
